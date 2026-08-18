# _src/*.md 4종을 한국어 reference.docx 적용해 docx로 일괄 변환 + 재오픈 검증
import subprocess, sys, os, glob
from docx import Document

SRC = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.dirname(SRC)
REF = os.path.join(SRC, "reference.docx")

ORDER = ["00_회의핵심안건", "01_회의진행자료", "02_위험도심층분석", "03_데이터인사이트", "04_다음스텝권고"]

def mojibake(t):
    return ('�' in t) or ('Ã' in t and 'á' in t)

ok = True
for stem in ORDER:
    md = os.path.join(SRC, stem + ".md")
    if not os.path.exists(md):
        print(f"[누락] {stem}.md 없음"); ok = False; continue
    docx = os.path.join(OUT, stem + ".docx")
    r = subprocess.run(["pandoc", md, "-o", docx, "--reference-doc", REF, "--toc", "--toc-depth=2"],
                       capture_output=True, text=True)
    if r.returncode != 0:
        print(f"[실패] {stem}: {r.stderr[:300]}"); ok = False; continue
    d = Document(docx)
    txt = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells: txt += "\n" + c.text
    has_ko = any('가' <= ch <= '힣' for ch in txt)
    print(f"[OK] {stem}.docx | paras={len(d.paragraphs)} tables={len(d.tables)} 한글={has_ko} 모지바케={mojibake(txt)} 길이={len(txt)}")
    if not has_ko or mojibake(txt): ok = False

print("=== 전체:", "통과" if ok else "문제있음", "===")
sys.exit(0 if ok else 1)
